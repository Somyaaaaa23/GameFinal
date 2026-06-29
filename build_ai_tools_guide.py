from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


OUTPUT = "/Users/sagaragarwal/GameFinal/Under-the-Radar_AI_Tools_Free_Tiers_2026.docx"

NAVY = "15324B"
BLUE = "247BA0"
TEAL = "2A9D8F"
GOLD = "E9C46A"
INK = "24313A"
MUTED = "66737D"
PALE = "EEF5F7"
LIGHT = "F6F8FA"
WHITE = "FFFFFF"
RED = "9C3D3D"


tools = [
    {
        "name": "Napkin AI",
        "category": "Visual communication",
        "url": "https://www.napkin.ai/pricing/",
        "best": "Turning text into clean diagrams, process visuals, and presentation-ready explainers.",
        "free": "Free forever; 500 AI credits per week, unlimited visual editing and file imports, plus PNG and PDF exports.",
        "catch": "Free exports carry Napkin branding; PPT and SVG export require a paid plan.",
        "why": "It removes the blank-canvas problem and is unusually useful for reports, decks, lessons, and product documentation.",
        "try": "Paste a dense paragraph from a report and generate three visual explanations.",
        "badge": "BEST QUICK WIN",
    },
    {
        "name": "Elicit",
        "category": "Research",
        "url": "https://elicit.com/pricing",
        "best": "Literature reviews, finding papers, comparing evidence, and chatting with full-text research.",
        "free": "Basic plan is free, with two automated reports per month, broad paper search, unlimited summaries, and source visibility.",
        "catch": "Advanced systematic-review workflows, larger extractions, exports, and API access are paid.",
        "why": "Its workflow is built around papers and evidence instead of generic web answers.",
        "try": "Ask a narrow research question, then compare the claims and study designs of five papers.",
        "badge": "RESEARCH PICK",
    },
    {
        "name": "Consensus",
        "category": "Research",
        "url": "https://consensus.app/pricing/",
        "best": "Getting evidence-led answers sourced from peer-reviewed research.",
        "free": "A free individual experience is available with limited advanced analyses and searches.",
        "catch": "Usage limits and feature allocations can change; verify the live pricing page before relying on a specific quota.",
        "why": "It is faster than manually translating a general question into academic search terms.",
        "try": "Test a debatable health or productivity claim and inspect the cited studies, not just the summary.",
        "badge": "FAST EVIDENCE",
    },
    {
        "name": "Krea",
        "category": "Images, video & creative",
        "url": "https://www.krea.ai/pricing",
        "best": "Real-time image generation, enhancement, editing, and experimentation across many creative models.",
        "free": "100 compute units per day, no credit card required; full real-time model access and limited image, video, 3D, lip-sync, upscale, and LoRA access.",
        "catch": "The free plan offers limited access to heavier models and workflows; commercial licensing is listed with paid plans.",
        "why": "Real-time visual feedback makes it excellent for exploring composition and style quickly.",
        "try": "Sketch a rough layout and iterate live until the composition works.",
        "badge": "CREATIVE PICK",
    },
    {
        "name": "Recraft",
        "category": "Images, video & creative",
        "url": "https://www.recraft.ai/pricing",
        "best": "Brand-style graphics, icons, vector-like artwork, image sets, and design variations.",
        "free": "A free plan provides generation credits and access to the core creative workspace.",
        "catch": "Free-plan images are public and owned by Recraft; commercial use and ownership have limitations.",
        "why": "It is more design-system-oriented than many prompt-to-image tools.",
        "try": "Generate a coordinated set of icons in one visual style rather than a single standalone image.",
        "badge": "DESIGN SYSTEMS",
    },
    {
        "name": "Granola",
        "category": "Meetings & notes",
        "url": "https://www.granola.ai/pricing",
        "best": "Enhancing your own meeting notes with transcripts and AI summaries without adding a visible meeting bot.",
        "free": "Basic is $0 and includes AI meeting notes, limited meeting history, cross-meeting chat, shared folders, templates, and multiple languages.",
        "catch": "Unlimited history, advanced models, integrations, API, and administration are paid.",
        "why": "The human-plus-AI note-taking workflow often produces more useful notes than a fully automatic recorder.",
        "try": "Write only decisions and unclear points during a meeting, then let Granola expand the notes.",
        "badge": "MEETING PICK",
    },
    {
        "name": "Tactiq",
        "category": "Meetings & notes",
        "url": "https://tactiq.io/pricing",
        "best": "Live transcription and AI actions for Google Meet, Zoom, and Microsoft Teams.",
        "free": "A free plan is available for light transcription and AI usage.",
        "catch": "Monthly transcription and AI-credit limits apply; check the current pricing page for exact allowances.",
        "why": "It lives close to the browser meeting workflow and makes quotes, tasks, and follow-ups easy to capture.",
        "try": "Create a reusable prompt that converts a transcript into decisions, owners, and deadlines.",
        "badge": "BROWSER MEETINGS",
    },
    {
        "name": "Brisk Teaching",
        "category": "Education",
        "url": "https://www.briskteaching.com/",
        "best": "Creating lessons, quizzes, feedback, reading-level adjustments, and instructional materials inside common education tools.",
        "free": "Core educator features are offered free; schools can purchase broader premium and administrative capabilities.",
        "catch": "Institution-level features and some advanced capabilities may require a paid school plan.",
        "why": "It works inside the pages and documents teachers already use rather than forcing a separate workflow.",
        "try": "Turn one article into three reading levels plus a short comprehension check.",
        "badge": "EDUCATOR PICK",
    },
    {
        "name": "Pieces",
        "category": "Developer productivity",
        "url": "https://pieces.app/",
        "best": "Remembering code, tabs, documents, and work context across desktop apps and developer tools.",
        "free": "The desktop product can be started free and supports local, on-device context and optional local or cloud models.",
        "catch": "It is most valuable after it has accumulated context; review capture settings before enabling broad activity memory.",
        "why": "Its long-term work memory tackles the repeated 'where did I see that?' problem.",
        "try": "Use natural-language search to recover a code snippet or page from work you did several days ago.",
        "badge": "DEV MEMORY",
    },
    {
        "name": "AnythingLLM",
        "category": "Private & local AI",
        "url": "https://anythingllm.com/",
        "best": "Chatting privately with documents, running local agents, and building a personal knowledge workspace.",
        "free": "Desktop is free, open source, MIT licensed, local by default, and requires no account.",
        "catch": "Local model quality and speed depend on your hardware; paid API models still incur their own provider costs.",
        "why": "It packages document chat, retrieval, models, storage, and agents into an approachable desktop app.",
        "try": "Create a workspace for a folder of PDFs and ask questions that require evidence from several files.",
        "badge": "PRIVATE DOCS",
    },
    {
        "name": "Jan",
        "category": "Private & local AI",
        "url": "https://www.jan.ai/",
        "best": "Running an open-source ChatGPT-style app with local models or optional online providers.",
        "free": "Free and open source, with downloads for desktop and support for open and online models.",
        "catch": "Large local models need substantial RAM and storage; online models may require paid API keys.",
        "why": "It gives non-specialists a clean route into local AI without building a stack from scratch.",
        "try": "Run a small local model offline, then compare its privacy, speed, and quality with a cloud model.",
        "badge": "LOCAL CHAT",
    },
    {
        "name": "LM Studio",
        "category": "Private & local AI",
        "url": "https://lmstudio.ai/",
        "best": "Discovering, downloading, testing, and serving local language models from a desktop interface.",
        "free": "Free for personal use, with local chat and an OpenAI-compatible local server.",
        "catch": "Model licenses differ, and performance depends heavily on memory, GPU support, and quantization.",
        "why": "It is one of the simplest ways to compare local models and expose them to other apps.",
        "try": "Load two small models and compare them on the same structured prompt.",
        "badge": "LOCAL MODEL LAB",
    },
    {
        "name": "Pinokio",
        "category": "Private & local AI",
        "url": "https://pinokio.computer/",
        "best": "Installing and running open-source AI apps through scripted, one-click workflows.",
        "free": "The launcher is free; many apps it installs are open source and run locally.",
        "catch": "Disk usage can grow quickly, and third-party scripts still deserve security scrutiny.",
        "why": "It lowers the command-line barrier for trying local AI apps.",
        "try": "Check disk and hardware requirements before installing.",
        "badge": "AI APP LAUNCHER",
    },
    {
        "name": "Upscayl",
        "category": "Images, video & creative",
        "url": "https://upscayl.org/",
        "best": "Improving the resolution of low-quality images locally with AI upscaling.",
        "free": "Free and open source for desktop use.",
        "catch": "Results depend on the source image and selected model; AI upscaling cannot recover truly missing detail.",
        "why": "It is focused, private, and useful for old assets, screenshots, and small generated images.",
        "try": "Compare two upscale models on faces, text, and illustration edges before choosing a result.",
        "badge": "FREE UTILITY",
    },
    {
        "name": "Hugging Face Spaces",
        "category": "Experimentation",
        "url": "https://huggingface.co/spaces",
        "best": "Trying thousands of community AI demos for text, image, audio, video, 3D, and multimodal tasks.",
        "free": "Many public Spaces can be used free, subject to queues, sleeping apps, and creator-set limits.",
        "catch": "Quality, privacy, uptime, and licensing vary by Space; never upload sensitive data without checking the app.",
        "why": "It is the best discovery layer for testing emerging open models before installing anything.",
        "try": "Search for one task, test three Spaces with the same input, and note model and license details.",
        "badge": "DISCOVERY HUB",
    },
]


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=None, bold=None, color=INK, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(c)
    rpr.append(u)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    link.append(r)
    paragraph._p.append(link)


def add_field(paragraph, field):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_label_line(doc, label, text, label_color=TEAL):
    p = doc.add_paragraph(style="Body Compact")
    r = p.add_run(label + " ")
    set_font(r, 9.2, True, label_color)
    r2 = p.add_run(text)
    set_font(r2, 9.2, False, INK)
    return p


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(0.82)
sec.right_margin = Inches(0.82)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(10.2)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.13

for name, size, color, before, after in [
    ("Title", 29, NAVY, 0, 8),
    ("Heading 1", 18, NAVY, 14, 7),
    ("Heading 2", 13, BLUE, 10, 4),
    ("Heading 3", 10.5, TEAL, 7, 2),
]:
    st = styles[name]
    st.font.name = "Aptos Display" if name != "Heading 3" else "Aptos"
    st._element.rPr.rFonts.set(qn("w:ascii"), st.font.name)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), st.font.name)
    st.font.size = Pt(size)
    st.font.bold = name != "Title"
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

body_compact = styles.add_style("Body Compact", WD_STYLE_TYPE.PARAGRAPH)
body_compact.font.name = "Aptos"
body_compact._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
body_compact._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
body_compact.font.size = Pt(9.2)
body_compact.font.color.rgb = RGBColor.from_string(INK)
body_compact.paragraph_format.space_after = Pt(3)
body_compact.paragraph_format.line_spacing = 1.08

small = styles.add_style("Small Muted", WD_STYLE_TYPE.PARAGRAPH)
small.font.name = "Aptos"
small.font.size = Pt(8.2)
small.font.color.rgb = RGBColor.from_string(MUTED)
small.paragraph_format.space_after = Pt(3)
small.paragraph_format.line_spacing = 1.05

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = header.add_run("FIELD GUIDE  |  JUNE 2026")
set_font(r, 8, True, MUTED)

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("Under-the-Radar AI Tools  •  ")
set_font(r, 8, False, MUTED)
add_field(footer, "PAGE")

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(46)
p.paragraph_format.space_after = Pt(12)
r = p.add_run("2026 FIELD GUIDE")
set_font(r, 10, True, TEAL)

p = doc.add_paragraph(style="Title")
p.paragraph_format.keep_with_next = True
r = p.add_run("Under-the-Radar\nAI Tools")
set_font(r, 29, True, NAVY, "Aptos Display")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
r = p.add_run("15 genuinely useful tools with free tiers or free local editions")
set_font(r, 14, False, BLUE, "Aptos Display")

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
widths = [Inches(2.15), Inches(2.15), Inches(2.15)]
for i, (cell, width) in enumerate(zip(table.rows[0].cells, widths)):
    cell.width = width
    set_cell_margins(cell, 150, 140, 150, 140)
    shade(cell, [NAVY, TEAL, GOLD][i])
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
for cell, label, value, color in zip(
    table.rows[0].cells,
    ["TOOLS", "CATEGORIES", "PRICE TO START"],
    ["15", "7", "$0"],
    [WHITE, WHITE, NAVY],
):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(value + "\n")
    set_font(r, 20, True, color, "Aptos Display")
    r = p.add_run(label)
    set_font(r, 7.5, True, color)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(22)
p.paragraph_format.space_after = Pt(7)
r = p.add_run("How this guide was selected")
set_font(r, 12, True, NAVY)

for text in [
    "Useful beyond a novelty demo, with a clear job it performs well.",
    "Less mainstream than the biggest general-purpose AI brands.",
    "A real free tier, free desktop edition, open-source release, or meaningful free access.",
    "Official product and pricing pages checked on June 13, 2026 where available.",
]:
    p = doc.add_paragraph(style="Body Compact")
    p.style = styles["Body Compact"]
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    r = p.add_run("●  ")
    set_font(r, 8, True, TEAL)
    r = p.add_run(text)
    set_font(r, 9.5, False, INK)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(13)
p.paragraph_format.space_after = Pt(0)
shade_color = PALE
p_pr = p._p.get_or_add_pPr()
shd = OxmlElement("w:shd")
shd.set(qn("w:fill"), shade_color)
p_pr.append(shd)
r = p.add_run("Important: ")
set_font(r, 9.2, True, RED)
r = p.add_run("Free plans change often. Recheck the linked official page before starting a critical workflow, uploading sensitive data, or assuming commercial-use rights.")
set_font(r, 9.2, False, INK)

doc.add_page_break()

# Quick picks
doc.add_heading("Start Here: The Fastest Picks", level=1)
p = doc.add_paragraph()
r = p.add_run("Choose by outcome, not by feature count.")
set_font(r, 11, False, MUTED)

quick = [
    ("Make an explanation visual", "Napkin AI", "Paste text and generate an editable diagram."),
    ("Research a serious question", "Elicit", "Search papers, compare evidence, and inspect sources."),
    ("Explore creative directions", "Krea", "Generate and edit visuals with real-time feedback."),
    ("Take better meeting notes", "Granola", "Add your own notes, then let AI enrich them."),
    ("Chat privately with documents", "AnythingLLM", "Build a local workspace around your files."),
    ("Run AI without the cloud", "Jan or LM Studio", "Download a small local model and test it offline."),
    ("Recover forgotten work context", "Pieces", "Search across code, tabs, and documents by meaning."),
    ("Discover emerging AI demos", "Hugging Face Spaces", "Test community models before installing them."),
]
qt = doc.add_table(rows=1, cols=3)
qt.alignment = WD_TABLE_ALIGNMENT.CENTER
qt.autofit = False
for c, w in zip(qt.rows[0].cells, [Inches(2.1), Inches(1.55), Inches(3.15)]):
    c.width = w
    shade(c, NAVY)
    set_cell_margins(c)
for c, text in zip(qt.rows[0].cells, ["I WANT TO…", "USE", "FIRST MOVE"]):
    p = c.paragraphs[0]
    r = p.add_run(text)
    set_font(r, 8.5, True, WHITE)
set_repeat_table_header(qt.rows[0])
for goal, tool, move in quick:
    cells = qt.add_row().cells
    for c, w in zip(cells, [Inches(2.1), Inches(1.55), Inches(3.15)]):
        c.width = w
        set_cell_margins(c, 110, 130, 110, 130)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade(cells[1], PALE)
    for c, text, bold, color in [
        (cells[0], goal, False, INK),
        (cells[1], tool, True, TEAL),
        (cells[2], move, False, INK),
    ]:
        p = c.paragraphs[0]
        r = p.add_run(text)
        set_font(r, 8.8, bold, color)

doc.add_heading("Free Does Not Mean Risk-Free", level=1)
for label, text in [
    ("Privacy", "Meeting, research, memory, and document tools may process sensitive material. Check retention, model-training, sharing, and deletion controls."),
    ("Commercial use", "Free image plans may make outputs public, retain ownership, or limit commercial rights. Recraft is a notable example."),
    ("Local hardware", "Local tools avoid many cloud privacy concerns, but models can consume several gigabytes of storage and significant RAM."),
    ("Reliability", "Community demos can sleep, queue, break, or disappear. Do not build a critical process around an unmaintained demo."),
    ("Accuracy", "Research assistants speed up discovery; they do not remove the need to read sources and assess study quality."),
]:
    add_label_line(doc, label.upper() + ":", text)

# Category sections
category_order = [
    "Research",
    "Visual communication",
    "Images, video & creative",
    "Meetings & notes",
    "Education",
    "Developer productivity",
    "Private & local AI",
    "Experimentation",
]

for category in category_order:
    selected = [t for t in tools if t["category"] == category]
    if not selected:
        continue
    heading = doc.add_heading(category, level=1)
    heading.paragraph_format.page_break_before = category != "Experimentation"
    intro = {
        "Research": "Use these to find and organize evidence. Treat the generated answer as a map to sources, not the final authority.",
        "Visual communication": "Use AI to convert ideas into diagrams and explainers that can be edited, exported, and reused.",
        "Images, video & creative": "The strongest free creative tools are excellent for exploration, but output rights and credit limits matter.",
        "Meetings & notes": "Choose based on whether you want a visible meeting bot, browser transcription, or an AI-enhanced personal notepad.",
        "Education": "Teacher-focused AI is most useful when it adapts existing material and stays inside familiar classroom workflows.",
        "Developer productivity": "Memory tools can be more valuable than another code generator because they preserve context across fragmented work.",
        "Private & local AI": "These tools run models or document workflows on your device. They trade cloud convenience for privacy and hardware demands.",
        "Experimentation": "Discovery platforms help you test new models quickly, but every community app has its own privacy, licensing, and reliability profile.",
    }[category]
    p = doc.add_paragraph()
    r = p.add_run(intro)
    set_font(r, 10.5, False, MUTED)

    for idx, tool in enumerate(selected):
        if idx:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(2)
            p_pr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "DCE5EA")
            pbdr.append(bottom)
            p_pr.append(pbdr)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(tool["name"])
        set_font(r, 15, True, NAVY, "Aptos Display")
        r = p.add_run("   " + tool["badge"])
        set_font(r, 7.5, True, TEAL)

        p = doc.add_paragraph(style="Body Compact")
        p.paragraph_format.keep_with_next = True
        hyperlink(p, "Official site / pricing", tool["url"])

        add_label_line(doc, "BEST FOR:", tool["best"])
        add_label_line(doc, "FREE ACCESS:", tool["free"])
        add_label_line(doc, "THE CATCH:", tool["catch"], RED)
        add_label_line(doc, "WHY IT STANDS OUT:", tool["why"])
        add_label_line(doc, "TRY THIS:", tool["try"], BLUE)

# Evaluation checklist
heading = doc.add_heading("A 10-Minute Evaluation Checklist", level=1)
heading.paragraph_format.page_break_before = True
p = doc.add_paragraph()
r = p.add_run("Before adopting any free AI tool, run one realistic task and score it on these six questions.")
set_font(r, 10.5, False, MUTED)

checks = [
    ("1", "Output quality", "Would you use the result after a normal amount of editing?"),
    ("2", "Workflow fit", "Does it reduce steps inside the tools you already use?"),
    ("3", "Free-tier honesty", "Is the limit clear, renewable, and sufficient for your real usage?"),
    ("4", "Data handling", "Can you explain where your inputs go, how long they stay, and who can access them?"),
    ("5", "Exit path", "Can you export your work in a useful format if the product changes?"),
    ("6", "Rights and licensing", "Can you legally use the output for your intended personal or commercial purpose?"),
]
for num, title, text in checks:
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    a, b = t.rows[0].cells
    a.width = Inches(0.55)
    b.width = Inches(6.25)
    shade(a, TEAL)
    shade(b, LIGHT)
    set_cell_margins(a, 110, 80, 110, 80)
    set_cell_margins(b, 110, 150, 110, 150)
    p = a.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(num)
    set_font(r, 12, True, WHITE)
    p = b.paragraphs[0]
    r = p.add_run(title + "  ")
    set_font(r, 10, True, NAVY)
    r = p.add_run(text)
    set_font(r, 9.2, False, INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)

doc.add_heading("Suggested Starter Stack", level=1)
stacks = [
    ("Student or researcher", "Elicit + Consensus + Napkin AI", "Find evidence, cross-check the claim, then explain it visually."),
    ("Creator or marketer", "Krea + Recraft + Upscayl", "Explore concepts, create coordinated assets, then improve final resolution."),
    ("Knowledge worker", "Granola + Pieces + Napkin AI", "Capture meetings, recover context, and communicate decisions."),
    ("Privacy-first user", "AnythingLLM + Jan or LM Studio", "Chat with documents and run models locally."),
    ("Teacher", "Brisk Teaching + Napkin AI", "Adapt learning material and add clear visual explanations."),
]
for audience, stack, workflow in stacks:
    p = doc.add_paragraph(style="Body Compact")
    r = p.add_run(audience + ": ")
    set_font(r, 9.4, True, NAVY)
    r = p.add_run(stack + ". ")
    set_font(r, 9.4, True, TEAL)
    r = p.add_run(workflow)
    set_font(r, 9.4, False, INK)

# Sources
heading = doc.add_heading("Official Sources", level=1)
heading.paragraph_format.page_break_before = True
p = doc.add_paragraph()
r = p.add_run("Access and pricing checked June 13, 2026. Exact quotas and rights can change without notice.")
set_font(r, 9.5, False, MUTED)

for t in tools:
    p = doc.add_paragraph(style="Small Muted")
    r = p.add_run(t["name"] + " — ")
    set_font(r, 8.5, True, NAVY)
    hyperlink(p, t["url"], t["url"], BLUE)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
r = p.add_run("Selection note")
set_font(r, 10.5, True, NAVY)
p = doc.add_paragraph(style="Small Muted")
r = p.add_run(
    "“Under-the-radar” is subjective. This guide intentionally excludes the largest general-purpose assistants and focuses on specialized tools that offer a useful free starting point. It is a curated field guide, not an exhaustive directory of every AI product."
)
set_font(r, 8.5, False, MUTED)

# Keep rows and set table borders/geometry.
for table in doc.tables:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05

doc.core_properties.title = "Under-the-Radar AI Tools with Free Tiers — 2026"
doc.core_properties.subject = "A curated field guide to useful specialized AI tools"
doc.core_properties.author = "OpenAI"
doc.core_properties.keywords = "AI tools, free tier, local AI, research, creative AI, productivity"
doc.save(OUTPUT)
print(OUTPUT)
